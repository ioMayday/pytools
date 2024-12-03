# -*- coding: utf-8 -*-
"""
Created on Tue Oct 29 22:20:16 2024

@author: Administrator

https://blog.csdn.net/m0_54985816/article/details/128588900

文件名：替换为1（完成制作款式1）、2（复习制作款式1）、3、（完成制作款式2）4、（复习制作款式2）



基线：1、2word


第1步：循环复制文件，i

https://geek-docs.com/python/python-ask-answer/454_python_how_to_use_pythondocx_to_replace_text_in_a_word_document_and_save.html
"""

from docx import Document


doc_path = r'C:\Users\Administrator\Desktop\1、宝安为明小学课后服务课程教案(创意美术 小学）.docx'
doc_new_path = r'C:\Users\Administrator\Desktop\new.docx'
str_src = '完成绘画'
str_tar_odd = '完成制作款式'
str_tar_even = '复习制作款式'

# doc = Document(r'C:\Users\Administrator\Desktop\1、宝安为明小学课后服务课程教案(创意美术 小学）.docx')  # 替换为你要操作的Word文档路径

# for paragraph in doc.paragraphs:
#     if str_src in paragraph.text:
#         paragraph.text = paragraph.text.replace(str_src, str_tar)
# doc.save('new_document.docx')  # 替换为你要保存的新文档的路径和名称


def read_ducment(old,new, document):
    # 遍历文档
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            #替换功能
            if old in run.text:
                run.text=run.text.replace(old,new)
 
    # 遍历表格
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                #遍历表格段落内容，回到上个步骤，将cell当作paragraph处理
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        #替换功能
                        if old in cell.text:
                            run.text=run.text.replace(old,new)


# # test
# document = Document(doc_path)
# read_ducment(str_src,str_tar, document)
# document.save(doc_new_path)


cnt = 74
doc_new_dir = r'C:\Users\Administrator\Desktop\tmp'
doc_path_origin = r'C:\Users\Administrator\Desktop\tmp\0、宝安为明小学课后服务课程教案(棉花糖 小学）.docx'

cnt_d2 = cnt // 2
str_split = '\\'

for i in range(cnt_d2):
    k = i + 1
    str_file_name = r'、宝安为明小学课后服务课程教案(棉花糖 小学）.docx'
    doc_new_path_odd = doc_new_dir  + str_split + str(2*k-1) + str_file_name
    str_tar_odd_conca = str_tar_odd + str(k)
    document_odd = Document(doc_path_origin)
    read_ducment(str_src, str_tar_odd_conca, document_odd)
    document_odd.save(doc_new_path_odd)
    
    doc_new_path_even = doc_new_dir + str_split + str(2*k) + str_file_name
    str_tar_even_conca = str_tar_even + str(k)
    document_even = Document(doc_path_origin)
    read_ducment(str_src, str_tar_even_conca, document_even)
    document_even.save(doc_new_path_even)
